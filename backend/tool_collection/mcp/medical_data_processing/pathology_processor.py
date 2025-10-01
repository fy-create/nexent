"""
病理学数据处理器
专门处理病理教科书等医疗专业素材
"""

import logging
import json
import re
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import asyncio
from dataclasses import dataclass

logger = logging.getLogger(__name__)

@dataclass
class PathologyDocument:
    """病理学文档数据结构"""
    title: str
    content: str
    chapter: str
    section: str
    medical_terms: List[str]
    diseases: List[str]
    histological_features: List[str]
    source_file: str
    page_number: Optional[int] = None

class PathologyDataProcessor:
    """病理学数据处理器"""
    
    def __init__(self):
        self.medical_term_patterns = self._load_medical_patterns()
        self.disease_patterns = self._load_disease_patterns()
        self.histology_patterns = self._load_histology_patterns()
        
    def _load_medical_patterns(self) -> List[str]:
        """加载医学术语识别模式"""
        return [
            r'[A-Za-z]+细胞',
            r'[A-Za-z]*瘤',
            r'[A-Za-z]*癌',
            r'[A-Za-z]*病',
            r'[A-Za-z]*症',
            r'病理[学性]?',
            r'组织[学性]?',
            r'形态[学性]?',
            r'免疫组化',
            r'HE染色',
            r'特殊染色',
            r'分子[标记物]*',
            r'诊断[标准]*',
            r'鉴别诊断',
            r'预后[因子]*'
        ]
    
    def _load_disease_patterns(self) -> List[str]:
        """加载疾病诊断模式"""
        return [
            r'恶性[肿瘤|肿块]*',
            r'良性[肿瘤|肿块]*',
            r'[A-Za-z]*腺癌',
            r'[A-Za-z]*鳞癌',
            r'[A-Za-z]*肉瘤',
            r'[A-Za-z]*淋巴瘤',
            r'[A-Za-z]*白血病',
            r'炎症[性反应]*',
            r'增生[性病变]*',
            r'化生[性改变]*',
            r'不典型增生',
            r'原位癌',
            r'浸润性癌'
        ]
    
    def _load_histology_patterns(self) -> List[str]:
        """加载组织学特征模式"""
        return [
            r'细胞[异型性|多形性]*',
            r'核[分裂象|质比]*',
            r'胞质[嗜酸性|嗜碱性]*',
            r'细胞[排列|结构]*',
            r'腺体[结构|形成]*',
            r'间质[反应|浸润]*',
            r'血管[侵犯|形成]*',
            r'神经[侵犯|浸润]*',
            r'坏死[区域|灶]*',
            r'纤维化',
            r'钙化',
            r'出血',
            r'水肿'
        ]
    
    async def process_pathology_textbook(self, 
                                       file_path: str,
                                       output_format: str = "structured") -> Dict[str, Any]:
        """
        处理病理教科书
        
        Args:
            file_path: 病理教科书文件路径
            output_format: 输出格式 (structured, qa_pairs, knowledge_graph)
            
        Returns:
            处理后的结构化数据
        """
        try:
            logger.info(f"开始处理病理教科书: {file_path}")
            
            # 1. 读取和解析文档
            raw_content = await self._read_document(file_path)
            
            # 2. 文档结构化分析
            structured_docs = await self._structure_document(raw_content, file_path)
            
            # 3. 医学术语提取和标注
            annotated_docs = await self._annotate_medical_content(structured_docs)
            
            # 4. 根据输出格式生成结果
            if output_format == "structured":
                result = await self._generate_structured_output(annotated_docs)
            elif output_format == "qa_pairs":
                result = await self._generate_qa_pairs(annotated_docs)
            elif output_format == "knowledge_graph":
                result = await self._generate_knowledge_graph(annotated_docs)
            else:
                raise ValueError(f"不支持的输出格式: {output_format}")
            
            logger.info(f"病理教科书处理完成，生成 {len(result.get('documents', []))} 个文档单元")
            return result
            
        except Exception as e:
            logger.error(f"处理病理教科书失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "processed_documents": 0
            }
    
    async def _read_document(self, file_path: str) -> str:
        """读取文档内容"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 根据文件类型选择读取方式
        if file_path.suffix.lower() == '.pdf':
            return await self._read_pdf(file_path)
        elif file_path.suffix.lower() in ['.txt', '.md']:
            return await self._read_text(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            return await self._read_word(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_path.suffix}")
    
    async def _read_pdf(self, file_path: Path) -> str:
        """读取PDF文件"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content = ""
                for page in reader.pages:
                    content += page.extract_text() + "\n"
                return content
        except ImportError:
            raise ImportError("需要安装PyPDF2库: pip install PyPDF2")
    
    async def _read_text(self, file_path: Path) -> str:
        """读取文本文件"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    async def _read_word(self, file_path: Path) -> str:
        """读取Word文件"""
        try:
            import docx
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            return content
        except ImportError:
            raise ImportError("需要安装python-docx库: pip install python-docx")
    
    async def _structure_document(self, content: str, source_file: str) -> List[PathologyDocument]:
        """结构化文档内容"""
        documents = []
        
        # 按章节分割内容
        chapters = self._split_by_chapters(content)
        
        for chapter_title, chapter_content in chapters.items():
            # 按小节分割
            sections = self._split_by_sections(chapter_content)
            
            for section_title, section_content in sections.items():
                # 创建文档对象
                doc = PathologyDocument(
                    title=section_title,
                    content=section_content,
                    chapter=chapter_title,
                    section=section_title,
                    medical_terms=[],
                    diseases=[],
                    histological_features=[],
                    source_file=source_file
                )
                documents.append(doc)
        
        return documents
    
    def _split_by_chapters(self, content: str) -> Dict[str, str]:
        """按章节分割内容"""
        # 识别章节标题模式
        chapter_pattern = r'^第[一二三四五六七八九十\d]+章\s*(.+?)$'
        chapters = {}
        current_chapter = "未分类章节"
        current_content = ""
        
        lines = content.split('\n')
        for line in lines:
            chapter_match = re.match(chapter_pattern, line.strip())
            if chapter_match:
                # 保存前一章节
                if current_content.strip():
                    chapters[current_chapter] = current_content.strip()
                # 开始新章节
                current_chapter = chapter_match.group(1)
                current_content = ""
            else:
                current_content += line + "\n"
        
        # 保存最后一章节
        if current_content.strip():
            chapters[current_chapter] = current_content.strip()
        
        return chapters
    
    def _split_by_sections(self, content: str) -> Dict[str, str]:
        """按小节分割内容"""
        # 识别小节标题模式
        section_pattern = r'^(\d+\.\d+\s*.+?)$'
        sections = {}
        current_section = "未分类小节"
        current_content = ""
        
        lines = content.split('\n')
        for line in lines:
            section_match = re.match(section_pattern, line.strip())
            if section_match:
                # 保存前一小节
                if current_content.strip():
                    sections[current_section] = current_content.strip()
                # 开始新小节
                current_section = section_match.group(1)
                current_content = ""
            else:
                current_content += line + "\n"
        
        # 保存最后一小节
        if current_content.strip():
            sections[current_section] = current_content.strip()
        
        return sections
    
    async def _annotate_medical_content(self, documents: List[PathologyDocument]) -> List[PathologyDocument]:
        """标注医学内容"""
        for doc in documents:
            # 提取医学术语
            doc.medical_terms = self._extract_medical_terms(doc.content)
            
            # 提取疾病诊断
            doc.diseases = self._extract_diseases(doc.content)
            
            # 提取组织学特征
            doc.histological_features = self._extract_histological_features(doc.content)
        
        return documents
    
    def _extract_medical_terms(self, content: str) -> List[str]:
        """提取医学术语"""
        terms = set()
        for pattern in self.medical_term_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            terms.update(matches)
        return list(terms)
    
    def _extract_diseases(self, content: str) -> List[str]:
        """提取疾病诊断"""
        diseases = set()
        for pattern in self.disease_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            diseases.update(matches)
        return list(diseases)
    
    def _extract_histological_features(self, content: str) -> List[str]:
        """提取组织学特征"""
        features = set()
        for pattern in self.histology_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            features.update(matches)
        return list(features)
    
    async def _generate_structured_output(self, documents: List[PathologyDocument]) -> Dict[str, Any]:
        """生成结构化输出"""
        return {
            "success": True,
            "processing_type": "structured",
            "total_documents": len(documents),
            "documents": [
                {
                    "title": doc.title,
                    "content": doc.content,
                    "chapter": doc.chapter,
                    "section": doc.section,
                    "medical_terms": doc.medical_terms,
                    "diseases": doc.diseases,
                    "histological_features": doc.histological_features,
                    "source_file": doc.source_file
                }
                for doc in documents
            ],
            "statistics": {
                "total_medical_terms": sum(len(doc.medical_terms) for doc in documents),
                "total_diseases": sum(len(doc.diseases) for doc in documents),
                "total_histological_features": sum(len(doc.histological_features) for doc in documents)
            }
        }
    
    async def _generate_qa_pairs(self, documents: List[PathologyDocument]) -> Dict[str, Any]:
        """生成Q&A对数据"""
        # 这里会调用QA生成器
        from .qa_generator import MedicalQAGenerator
        
        qa_generator = MedicalQAGenerator()
        qa_pairs = []
        
        for doc in documents:
            doc_qa_pairs = await qa_generator.generate_qa_from_document(doc)
            qa_pairs.extend(doc_qa_pairs)
        
        return {
            "success": True,
            "processing_type": "qa_pairs",
            "total_qa_pairs": len(qa_pairs),
            "qa_pairs": qa_pairs
        }
    
    async def _generate_knowledge_graph(self, documents: List[PathologyDocument]) -> Dict[str, Any]:
        """生成知识图谱"""
        # 构建知识图谱结构
        entities = set()
        relationships = []
        
        for doc in documents:
            # 添加实体
            entities.update(doc.medical_terms)
            entities.update(doc.diseases)
            entities.update(doc.histological_features)
            
            # 构建关系
            for disease in doc.diseases:
                for feature in doc.histological_features:
                    relationships.append({
                        "source": disease,
                        "relation": "has_feature",
                        "target": feature,
                        "context": doc.title
                    })
        
        return {
            "success": True,
            "processing_type": "knowledge_graph",
            "entities": list(entities),
            "relationships": relationships,
            "total_entities": len(entities),
            "total_relationships": len(relationships)
        }